#!/usr/bin/env python3
"""Build an ATS-friendly resume.docx for the Virtual Networx Data Analyst & Power BI Engineer role."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume.docx")
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
for section in doc.sections:
    section.top_margin = section.bottom_margin = Pt(36)
    section.left_margin = section.right_margin = Pt(54)

def center(text, size=10.5, bold=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); p.paragraph_format.space_after = Pt(2); return p

def name_line(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); p.paragraph_format.space_after = Pt(0)

def section_head(text):
    p = doc.add_paragraph(); r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x1F,0x38,0x64)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'1F3864')
    pbdr.append(bottom); pPr.append(pbdr)

def role(title, org_dates):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run(title).bold = True; p.runs[0].font.size = Pt(10.5)
    p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(2); r2 = p2.add_run(org_dates); r2.italic = True; r2.font.size = Pt(10)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    for i, seg in enumerate(text.split("**")):
        run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(10)

def para(text, size=10):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    for i, seg in enumerate(text.split("**")):
        run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(size)

name_line("Andrew Green")
center("Data Analyst & Power BI Engineer — Power BI · Microsoft Fabric · SQL", 10.5, bold=True)
center("Fort Lauderdale, FL (Remote, ET)  |  561.339.1806  |  andrewgreen04@gmail.com  |  linkedin.com/in/agreen8", 9.5)

section_head("Summary")
para("Power BI developer with 10+ years turning complex data into clear, decision-driving dashboards on the Microsoft stack. Deep Power BI (semantic modeling, DAX, Power Query) plus production Microsoft Fabric — lakehouses, Direct Lake semantic models, and pipelines that most Power BI developers have only seen in demos. Strong SQL, data modeling, and ETL, with a track record of stakeholder storytelling and rigorous data-quality validation. Currently on a remote consulting engagement (contract extended four times — the most of any consultant on the project).")

section_head("Technical Skills")
for line in [
 "**BI & Visualization:** Power BI (semantic models, DAX, Power Query / M), Microsoft Fabric (Lakehouse, Direct Lake, Pipelines, Dataflows, Golden Datasets), Deneb / Vega-Lite, SSRS / Power BI Report Builder, Excel / PowerPivot",
 "**Data & Query:** SQL, T-SQL, PL/SQL, dimensional modeling, star schema, semantic layer, reusable DAX measures",
 "**ETL / Orchestration:** Azure Data Factory, SSIS, Azure Synapse, incremental / watermark loaders, Python (pandas)",
 "**Data Quality:** profiling, reconciliation, EXCEPT-keyed validation, KPI definitions, documentation, UAT",
 "**Cloud:** Microsoft Azure, Azure SQL",
 "**Certifications (in progress):** PL-300 Power BI Data Analyst · DP-700 Fabric Data Engineer · DP-203 Azure Data Engineer",
]:
    bullet(line)

section_head("Experience")
role("Business Intelligence Developer — Power BI Specialist (Contract/Consultant)", "Carnival Cruise Line — Remote  |  Jan 2025–Present")
for b in [
 "Design and maintain multi-page **Power BI** dashboards and reports for fleet HVAC energy performance across 79 ships and 8 brands — KPI strips, per-ship trend small-multiples, worst-offender rankings, and a sensor drill-through diagnostics page.",
 "Early production adopter of **Microsoft Fabric** — built Lakehouse ingestion pipelines and a **Direct Lake semantic model** delivering near-real-time dashboards without dataset-refresh overhead.",
 "Build scalable Power BI semantic models with five calculation groups and reusable DAX measures against live Azure SQL, serving engineering, operations, and executive audiences.",
 "Partner with Engineering, Operations, and executive leadership to gather requirements and translate complex data into action-oriented dashboard narratives — iterating enhancements that improved stakeholder adoption by 35%.",
 "Enforce data accuracy and integrity: profiling and reconciliation across raw, transformed, and Power BI layers; caught a fleet-wide sentinel-value data-poison bug and added plausibility guards; validated refactors via EXCEPT-keyed diffs across multi-million-row views. Optimized SQL processing time by 25%.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer (Power BI)", "Basic Fun — Boca Raton, FL  |  Mar 2023–Jan 2025")
for b in [
 "Built and maintained enterprise **Power BI** dashboards covering sales, budget, forecast accuracy, inventory, and variance-to-plan — contributing to a 20% improvement in Performance-to-Plan.",
 "Built **Microsoft Fabric** lakehouses and published **golden datasets** — governed, reusable data products decoupling reporting from raw source dependencies.",
 "Designed the company's first enterprise data warehouse end to end — ingestion, ETL, dimensional modeling, star schema, semantic layer, and Power BI reporting.",
 "Partnered with demand planning on predictive inventory modeling, achieving a 25% YoY improvement in forecast accuracy.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer", "Twin-Star International — Delray Beach, FL  |  Oct 2017–Mar 2023")
for b in [
 "Led enterprise Power BI deployment across the organization; defined the reporting/analytics roadmap.",
 "Architected the company's first enterprise data warehouse: dimensional models, star schema, SSIS ETL, and Azure Synapse for analytical processing.",
 "Automated P&L, Balance Sheet, and Cash Flow reporting across three portfolio companies on a unified chart of accounts — saving Finance 20–25 hours/month.",
 "Built paginated reports in SSRS / Power BI Report Builder for complex Excel and CSV data exports.",
]:
    bullet(b)

role("Business Analyst", "AutoNation — Palm Beach Gardens, FL  |  Apr 2016–Oct 2017")
for b in [
 "Created advanced DAX measures for real-time KPI tracking across financial forecasting, sales, and inventory.",
 "Designed SQL-based sales analytics and pricing tools supporting inventory optimization.",
]:
    bullet(b)

role("Financial Advisor", "Merrill Lynch — Palm Beach Gardens, FL  |  May 2009–Mar 2016")
bullet("Co-managed $250M in assets; applied ETL methodologies to aggregate regulatory data and ensure compliance.")

section_head("Education")
para("**B.S., Finance & Economics** — Florida State University, Tallahassee, FL")

doc.save(OUT); print("SAVED:", OUT)
