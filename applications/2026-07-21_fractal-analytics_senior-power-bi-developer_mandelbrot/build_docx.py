#!/usr/bin/env python3
"""Build an ATS-friendly resume.docx for the Fractal Analytics Senior Power BI Developer role."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume.docx")
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(10)
for section in doc.sections:
    section.top_margin = section.bottom_margin = Pt(20)
    section.left_margin = section.right_margin = Pt(42)

def center(text, size=10.5, bold=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); p.paragraph_format.space_after = Pt(2); return p

def name_line(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18); p.paragraph_format.space_after = Pt(0)

def section_head(text):
    p = doc.add_paragraph(); r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x1F,0x38,0x64)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'1F3864')
    pbdr.append(bottom); pPr.append(pbdr)

def role(title, org_dates):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
    p.add_run(title).bold = True; p.runs[0].font.size = Pt(10)
    p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(1); r2 = p2.add_run(org_dates); r2.italic = True; r2.font.size = Pt(9.5)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(1)
    for i, seg in enumerate(text.split("**")):
        run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(9.5)

def para(text, size=9.5):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    for i, seg in enumerate(text.split("**")):
        run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(size)

name_line("Andrew Green")
center("Senior Power BI Developer — Power BI · DAX · Azure Data Factory · Microsoft Fabric", 10.5, bold=True)
center("Fort Lauderdale, FL (Remote, ET)  |  561.339.1806  |  andrewgreen04@gmail.com  |  linkedin.com/in/agreen8", 9.5)

section_head("Summary")
para("Senior Power BI Developer with 10+ years on the Microsoft data stack — interactive dashboards, data models, DAX, and Power Query backed by Azure Data Factory, Synapse, and Data Lake pipelines. Built three enterprise analytics platforms from scratch across CPG and financial domains; production Microsoft Fabric adopter (Lakehouse, Direct Lake). Known for data accuracy (99%), performance optimization (25% faster processing), and stakeholder partnership that lifts adoption (35%).")

section_head("Technical Skills")
for line in [
 "**Power BI:** interactive dashboards & reports, semantic/data modeling, DAX (calculation groups), Power Query / M, SSRS / Report Builder, KPI-driven dashboard design",
 "**Azure:** Data Factory, Synapse Analytics, Data Lake, Azure SQL, SQL Server, Microsoft Fabric (Lakehouse, Direct Lake, OneLake, REST API / TMDL)",
 "**Data engineering:** SQL / T-SQL, ETL pipelines, SSIS, data warehousing, star-schema modeling, incremental loads & scheduled refresh, performance tuning, Python (pandas, scikit-learn), PL/SQL",
 "**Deployment & governance:** PBIP source control (Git), Fabric workspace deployment (REST API / TMDL), data profiling & reconciliation, UAT",
 "**Certifications (in progress):** PL-300 Power BI Data Analyst · DP-700 Fabric Data Engineer · DP-203 Azure Data Engineer",
]:
    bullet(line)

section_head("Experience")
role("Business Intelligence Developer — Power BI Specialist (Contract/Consultant)", "Carnival Cruise Line — Remote  |  Jan 2025–Present")
for b in [
 "Design and develop interactive **Power BI** dashboards and reports for fleet HVAC energy performance across 79 ships and 8 brands — from fleet KPI strips and ship rankings to a sensor-level drill-through diagnostics page.",
 "Create **data models, DAX measures, and Power Query transformations** supporting business logic — five calculation groups against live Azure SQL — while gathering requirements and **training business users** in working sessions that improved stakeholder adoption 35%.",
 "Integrate Power BI with Azure: three parallel **ETL pipelines** (IoT, cloud BMS, on-prem SQL) orchestrated in **Azure Synapse** with **incremental loads and scheduled refresh jobs** — 99% data accuracy for near-real-time reporting.",
 "Early production adopter of **Microsoft Fabric** — Lakehouse ingestion pipelines over **Azure Data Lake** (OneLake) and a **Direct Lake semantic model** delivering near-real-time dashboards without dataset-refresh overhead.",
 "Ensure **data accuracy, consistency, and performance optimization** across reports — 25% faster processing, ~16x columnstore compression, EXCEPT-keyed validation on multi-million-row views.",

]:
    bullet(b)

role("Senior Business Intelligence Engineer (Power BI)", "Basic Fun (consumer products / toys — CPG) — Boca Raton, FL  |  Mar 2023–Jan 2025")
for b in [
 "Designed the company's first enterprise **data warehouse** end to end — ingestion, ETL, star schema, semantic layer, Power BI — plus **Microsoft Fabric** lakehouses publishing governed **golden datasets**.",
 "Built **Azure Data Lake** ETL pipelines ingesting third-party point-of-sale data from web portal, FTP, email, and flat-file sources; implemented **incremental refresh** strategies replacing full loads, cutting pipeline run times.",
 "Built executive Power BI dashboards covering sales, budget, forecast accuracy, inventory, and variance-to-plan — contributing to a 20% improvement in Performance-to-Plan; automated **scheduled report distribution** via Power Automate.",

 "Partnered with demand planning on predictive inventory modeling, achieving a 25% YoY improvement in forecast accuracy.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer", "Twin-Star International (consumer products, 3-company portfolio) — Delray Beach, FL  |  Oct 2017–Mar 2023")
for b in [
 "Architected the company's first enterprise **data warehouse**: dimensional models, star schema, **SSIS** ETL packages, and Azure Synapse for analytical processing.",

 "Automated P&L, Balance Sheet, and Cash Flow reporting across three portfolio companies — saving Finance 20–25 hours/month.",
 "Migrated on-premise **SQL Server** schemas and siloed ERP data to **Azure SQL Server** via ETL pipelines.",
]:
    bullet(b)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
for i, seg in enumerate("**Earlier:** Business Analyst, AutoNation (2016-2017) - advanced DAX measures and SQL-based sales & pricing analytics  |  Financial Advisor, Merrill Lynch (2009-2016) - co-managed $250M in assets.  **Education:** B.S., Finance & Economics, Florida State University.".split("**")):
    run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(9.5)


doc.save(OUT); print("SAVED:", OUT)
