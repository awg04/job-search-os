#!/usr/bin/env python3
"""Build an ATS-friendly resume.docx for the Talenzaa Business Analyst (BI) role."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume.docx")
doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Pt(36); s.left_margin = s.right_margin = Pt(54)

def center(t, size=10.5, bold=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(size); p.paragraph_format.space_after = Pt(2)

def name_line(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = True; r.font.size = Pt(20); p.paragraph_format.space_after = Pt(0)

def section_head(t):
    p = doc.add_paragraph(); r = p.add_run(t.upper()); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x1F,0x38,0x64)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr(); b = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'6'); bt.set(qn('w:space'),'1'); bt.set(qn('w:color'),'1F3864'); b.append(bt); pPr.append(b)

def role(title, org):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run(title).bold = True; p.runs[0].font.size = Pt(10.5)
    p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(2); r = p2.add_run(org); r.italic = True; r.font.size = Pt(10)

def bullet(t):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    for i, seg in enumerate(t.split("**")):
        r = p.add_run(seg); r.bold = (i % 2 == 1); r.font.size = Pt(10)

def para(t, size=10):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    for i, seg in enumerate(t.split("**")):
        r = p.add_run(seg); r.bold = (i % 2 == 1); r.font.size = Pt(size)

name_line("Andrew Green")
center("Business Analyst — Business Intelligence · Power BI · Microsoft Azure", 10.5, bold=True)
center("Fort Lauderdale, FL (Remote, ET)  |  561.339.1806  |  andrewgreen04@gmail.com  |  linkedin.com/in/agreen8", 9.5)

section_head("Summary")
para("Business-intelligence professional with 10+ years turning business requirements into Power BI reporting and Azure data solutions. Strong on the full cycle — gathering and analyzing requirements, translating business needs into technical designs, and delivering dashboards executives act on. Deep Power BI plus Microsoft Azure data services (Azure Data Factory, Azure SQL, Synapse), SQL, data modeling, and ETL. Direct supply-chain / demand-planning domain experience (S&OP, forecasting). Currently on a remote consulting engagement (contract extended four times — the most of any consultant on the project).")

section_head("Technical Skills")
for line in [
 "**Business Analysis:** requirements gathering, source-to-target mapping, stakeholder management, data validation / UAT, documentation, Agile / Scrum",
 "**BI & Reporting:** Power BI (semantic models, DAX, Power Query), interactive dashboards, data visualization, SSRS / Power BI Report Builder, Microsoft Fabric",
 "**Azure Data Services:** Azure Data Factory (ADF), Azure SQL, Azure Synapse Analytics, data ingestion / storage / integration",
 "**Data:** SQL, T-SQL, PL/SQL, dimensional modeling, star schema, data warehousing, ETL (SSIS)",
 "**Domain:** supply chain / S&OP, demand & inventory planning, financial reporting (P&L, BS, CF)",
 "**Certifications (in progress):** PL-300 Power BI Data Analyst · DP-203 Azure Data Engineer · DP-700 Fabric Data Engineer",
]:
    bullet(line)

section_head("Experience")
role("Business Intelligence Developer — Power BI Specialist (Contract/Consultant)", "Carnival Cruise Line — Remote  |  Jan 2025–Present")
for b in [
 "Partner with Engineering, Operations, and executive stakeholders to **gather requirements**, define KPI logic, and translate business questions into Power BI dashboards — iterating enhancements that improved stakeholder adoption by 35%.",
 "Design multi-page Power BI reports and semantic models (DAX, calculation groups) against **Azure SQL**, serving operational and executive audiences.",
 "Orchestrate ETL with **Azure Synapse** across three source pipelines; validate data accuracy via reconciliation across raw, transformed, and Power BI layers. Reduced SQL processing time by 25%.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer (Power BI)", "Basic Fun — Boca Raton, FL  |  Mar 2023–Jan 2025")
for b in [
 "Built enterprise Power BI dashboards for sales, budget, forecast accuracy, inventory, and variance-to-plan — contributing to a 20% improvement in Performance-to-Plan.",
 "Partnered with **Demand Planning** on predictive inventory modeling, achieving a **25% YoY improvement in forecast accuracy** — direct supply-chain analytics.",
 "Designed the company's first enterprise data warehouse end to end — Azure data lake ingestion, ETL, dimensional modeling, star schema, semantic layer, and Power BI.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer", "Twin-Star International — Delray Beach, FL  |  Oct 2017–Mar 2023")
for b in [
 "Collaborated with **Sales & Operations Planning (S&OP)** to support demand planning, supply planning, and global supply-chain decisions.",
 "Architected the company's first enterprise data warehouse: dimensional models, star schema, **SSIS ETL**, and **Azure Synapse** for analytical processing; migrated on-prem SQL to **Azure SQL**.",
 "Led enterprise Power BI deployment; automated P&L, Balance Sheet, and Cash Flow reporting across three portfolio companies — saving Finance 20–25 hours/month.",
]:
    bullet(b)

role("Business Analyst", "AutoNation — Palm Beach Gardens, FL  |  Apr 2016–Oct 2017")
for b in [
 "Gathered requirements and produced daily/weekly cross-functional reports for Sales, Marketing, and Operations, improving operational visibility.",
 "Created advanced DAX measures for real-time KPI tracking across forecasting, sales, and inventory; built SQL-based analytics and pricing tools.",
]:
    bullet(b)

role("Financial Advisor", "Merrill Lynch — Palm Beach Gardens, FL  |  May 2009–Mar 2016")
bullet("Co-managed $250M in assets; applied ETL methodologies to aggregate regulatory data and ensure compliance.")

section_head("Education")
para("**B.S., Finance & Economics** — Florida State University, Tallahassee, FL")
doc.save(OUT); print("SAVED:", OUT)
