#!/usr/bin/env python3
"""Build an ATS-friendly resume.docx for the Armanino Senior BI & Analytics Developer role (JR103988)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume.docx")
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(10)
for section in doc.sections:
    section.top_margin = section.bottom_margin = Pt(16)
    section.left_margin = section.right_margin = Pt(40)

def center(text, size=10.5, bold=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); p.paragraph_format.space_after = Pt(2); return p

def name_line(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18); p.paragraph_format.space_after = Pt(0)

def section_head(text):
    p = doc.add_paragraph(); r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x1F,0x38,0x64)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'1F3864')
    pbdr.append(bottom); pPr.append(pbdr)

def role(title, org_dates):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
    p.add_run(title).bold = True; p.runs[0].font.size = Pt(10)
    p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(1); r2 = p2.add_run(org_dates); r2.italic = True; r2.font.size = Pt(9.5)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(1)
    for i, seg in enumerate(text.split("**")):
        run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(9.5)

def para(text, size=9.5):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    for i, seg in enumerate(text.split("**")):
        run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(size)

name_line("Andrew Green")
center("Senior BI & Analytics Developer — Microsoft Fabric · Power BI Semantic Modeling · SQL · Python", 10.5, bold=True)
center("Fort Lauderdale, FL (Boca-local / hybrid, ET)  |  561.339.1806  |  andrewgreen04@gmail.com  |  linkedin.com/in/agreen8", 9.5)

section_head("Summary")
para("Senior BI & Analytics Developer, 10+ years on the Microsoft data stack, working the seam between back-end data engineering and front-end reporting — building curated data models in SQL and Python across Microsoft Fabric (Lakehouse, notebooks, pipelines) and the Power BI semantic models on top. Three enterprise data warehouses built from scratch; production Fabric adopter with an FP&A/finance background. Known for 99% data accuracy, 25% faster processing, and Git-governed delivery.")

section_head("Technical Skills")
for line in [
 "**Microsoft Fabric:** Lakehouse & Warehouse, notebooks (transformation / enrichment / business logic), pipelines, Direct Lake, golden datasets, OneLake, REST API / TMDL",
 "**Power BI:** semantic / data modeling, DAX (calculation groups), Power Query / M, reusable models & clear metric definitions, SSRS",
 "**SQL & data engineering:** T-SQL (advanced querying, transformations, performance tuning), ETL, SSIS, Azure Data Factory & Synapse, incremental / watermark loads, refresh sequencing",
 "**Python & modeling:** Python (pandas, scikit-learn) for data transformation, dimensional modeling (star schema, fact / dimension, SCD), data validation & reconciliation",
 "**Governance & delivery:** Git / PBIP source control, peer review, naming & modeling standards, documentation, UAT",
 "**Certifications (in progress):** PL-300 Power BI Data Analyst · DP-700 Fabric Data Engineer · DP-203 Azure Data Engineer",
]:
    bullet(line)

section_head("Experience")
role("Business Intelligence Developer — Power BI Specialist (Contract/Consultant)", "Carnival Cruise Line — Remote  |  Jan 2025–Present")
for b in [
 "Bridge back-end data engineering and front-end reporting for the fleet HVAC energy platform — building curated data models in **SQL and Python** and the **Power BI semantic models** on top, across 79 ships and 8 brands.",
 "Build and maintain **Microsoft Fabric** notebooks and Lakehouse pipelines for transformation, enrichment, and business logic — ingesting IoT, cloud BMS, and on-prem SQL into a governed layer, with a **Direct Lake semantic model** for near-real-time reporting.",
 "Design semantic models with five **calculation groups**, reusable **DAX** measures, and clear, transparent **metric definitions** — enabling self-service and shifting stakeholders from static reporting to insight.",
 "Validate data outputs and business-rule alignment before models are surfaced — EXCEPT-keyed reconciliation across multi-million-row views — and own **refresh sequencing and dependencies** via staged loads in **Azure Synapse** (99% data accuracy).",
 "Work under **Git** (PBIP) with documented standards and naming conventions; tuned SQL for 25% faster processing and ~16x columnstore compression, testing changes before production.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer (Power BI)", "Basic Fun (consumer products / toys — CPG) — Boca Raton, FL  |  Mar 2023–Jan 2025")
for b in [
 "Designed the company's first enterprise **data warehouse** end to end — dimensional models, star schema, semantic layer, and Power BI — plus **Microsoft Fabric** lakehouses publishing governed **golden datasets** reusable across reports.",
 "Built **Azure Data Lake** and Fabric ETL/transformation pipelines ingesting third-party point-of-sale data from web portal, FTP, email, and flat-file sources; replaced full loads with **incremental refresh** to improve data readiness.",
 "Partnered with Finance, Sales, and Demand Planning to translate business questions into governed KPI models — executive dashboards on sales, budget, forecast, inventory, and variance-to-plan contributing to a 20% improvement in Performance-to-Plan.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer", "Twin-Star International (consumer products, 3-company portfolio) — Delray Beach, FL  |  Oct 2017–Mar 2023")
for b in [
 "Architected the company's first enterprise **data warehouse**: dimensional models, star schema, **SSIS** ETL packages, and Azure Synapse analytical processing.",
 "Automated **P&L, Balance Sheet, and Cash Flow** reporting across three portfolio companies on a unified chart of accounts — saving Finance 20–25 hours/month.",
 "Migrated on-premise **SQL Server** schemas and siloed ERP data to **Azure SQL** via ETL pipelines.",
]:
    bullet(b)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
for i, seg in enumerate("**Earlier:** Business Analyst, AutoNation (2016-2017) - advanced DAX measures and SQL-based sales & pricing analytics  |  Financial Advisor, Merrill Lynch (2009-2016) - co-managed $250M in assets.  **Education:** B.S., Finance & Economics, Florida State University.".split("**")):
    run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(9.5)

doc.save(OUT); print("SAVED:", OUT)
