#!/usr/bin/env python3
"""Build an ATS-friendly resume.docx for the Publicis Sapient Senior Associate Power BI Developer role."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume.docx")

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
for s in ("Heading 1", "Heading 2"):
    doc.styles[s].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

for section in doc.sections:
    section.top_margin = section.bottom_margin = Pt(36)
    section.left_margin = section.right_margin = Pt(54)

def name_line(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    return p

def center(text, size=10.5, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p

def section_head(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x1F,0x38,0x64)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'1F3864')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def role(title, org_dates):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title); r.bold = True; r.font.size = Pt(10.5)
    p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(org_dates); r2.italic = True; r2.font.size = Pt(10)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    # bold markers **x**
    parts = text.split("**")
    for i, seg in enumerate(parts):
        run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(10)

def para(text, size=10):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    parts = text.split("**")
    for i, seg in enumerate(parts):
        run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(size)

# Header
name_line("Andrew Green")
center("Senior Power BI Developer — Microsoft Power Platform (Power BI · Power Apps · Power Automate)", 10.5, bold=True)
center("Fort Lauderdale, FL (Remote, ET)  |  561.339.1806  |  andrewgreen04@gmail.com  |  linkedin.com/in/agreen8", 9.5)

# Summary
section_head("Summary")
para("Senior Power BI developer with 10+ years building enterprise reporting and analytics on the Microsoft stack. Deep Power BI — semantic modeling, DAX, and Power Query (M) — paired with hands-on Power Apps and Power Automate delivery: embedded write-back apps, automated distribution flows, and SharePoint / Microsoft 365 integration. Currently on a remote consulting engagement (contract extended four times — the most of any consultant on the project). Strong stakeholder partner who translates business requirements into governed, scalable technical deliverables.")

# Skills
section_head("Technical Skills")
for line in [
 "**Power Platform:** Power BI, Power Apps, Power Automate, Power BI Report Builder / SSRS, Power Query (M), DAX",
 "**Data Modeling:** Dimensional modeling, star schema, semantic models, calculation groups, reusable DAX measures",
 "**Integration:** Microsoft 365, SharePoint, Azure SQL, Azure Data Factory, Azure Synapse, REST APIs, external / 3rd-party data sources",
 "**Cloud & Platform:** Microsoft Azure, Microsoft Fabric (Lakehouse, Direct Lake, Pipelines, Golden Datasets)",
 "**Query / Languages:** SQL, T-SQL, PL/SQL, Python",
 "**Governance & Delivery:** Data validation & reconciliation, KPI definitions, documentation, UAT, Agile, Git (PBIP source control)",
 "**Certifications (in progress):** PL-300 Power BI Data Analyst · DP-700 Fabric Data Engineer · DP-203 Azure Data Engineer",
]:
    bullet(line)

# Experience
section_head("Experience")
role("Business Intelligence Developer — Power BI Specialist (Contract/Consultant)", "Carnival Cruise Line — Remote  |  Jan 2025–Present")
for b in [
 "Design and maintain multi-page Power BI dashboards and reports for fleet HVAC energy performance across 79 ships and 8 brands — KPI strips, per-ship trend small-multiples, worst-offender rankings, and a sensor drill-through diagnostics page.",
 "Build scalable Power BI semantic models with five calculation groups and reusable DAX measures serving engineering, operations, and executive audiences against live Azure SQL.",
 "Partner with Engineering, Operations, Fuel Performance, and executive leadership to gather requirements, define KPI logic, and translate complex data into action-oriented dashboards — iterating enhancements that improved stakeholder adoption by 35%.",
 "Optimized SQL extraction, transformation, and validation for large-scale datasets, reducing data processing time by 25% while improving refresh performance and reporting scalability.",
 "Established data-integrity practice: profiling, reconciliation across raw, transformed, and Power BI layers, and documentation of dataflows, KPI definitions, DAX measures, and refresh processes for governance and auditability.",
 "Early production adopter of Microsoft Fabric — built Lakehouse ingestion pipelines and a Direct Lake semantic model delivering near-real-time dashboards without dataset-refresh overhead. Report under Git (PBIP) source control.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer (Power BI)", "Basic Fun — Boca Raton, FL  |  Mar 2023–Jan 2025")
for b in [
 "Built and maintained enterprise Power BI dashboards covering sales, budget, forecast accuracy, inventory, and variance-to-plan — contributing to a 20% improvement in Performance-to-Plan outcomes.",
 "Automated daily/weekly report distributions and alerts with **Power Automate**, querying Power BI data on scheduled flows to improve stakeholder follow-through and reporting consistency.",
 "Designed the company's first enterprise data warehouse end to end — source ingestion, ETL, dimensional modeling, star schema, semantic layer, and Power BI reporting.",
 "Partnered with Finance, Sales, Demand Planning, and Operations to capture requirements, define KPI logic, and translate business questions into Power BI reports and DAX measures.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer", "Twin-Star International — Delray Beach, FL  |  Oct 2017–Mar 2023")
for b in [
 "Led enterprise Power BI deployment across the organization, defining the reporting/analytics roadmap and delivery accountability.",
 "Built a **Power Apps** write-back solution embedded in Power BI custom visuals, letting report users post commentary to a **SharePoint** Product Tracker without leaving the report — closing the loop between reporting and action tracking.",
 "Built paginated reports in SSRS / Power BI Report Builder for complex Excel and CSV data exports.",
 "Automated P&L, Balance Sheet, and Cash Flow reporting across three portfolio companies on a unified chart of accounts — saving Finance 20–25 hours/month.",
 "Architected the company's first enterprise data warehouse: dimensional models, star schema, SSIS ETL, and Azure Synapse for analytical processing.",
]:
    bullet(b)

role("Business Analyst", "AutoNation — Palm Beach Gardens, FL  |  Apr 2016–Oct 2017")
for b in [
 "Created advanced DAX measures for real-time KPI tracking across financial forecasting, sales performance, and inventory optimization.",
 "Designed SQL-based sales analytics and pricing tools supporting inventory optimization and bid-price calculation.",
]:
    bullet(b)

role("Financial Advisor", "Merrill Lynch — Palm Beach Gardens, FL  |  May 2009–Mar 2016")
bullet("Co-managed $250M in assets as junior partner on a wealth management team; applied ETL methodologies to aggregate regulatory data and ensure compliance with financial laws and internal policies — data governance in a regulated industry.")

# Education
section_head("Education")
para("**B.S., Finance & Economics** — Florida State University, Tallahassee, FL")

doc.save(OUT)
print("SAVED:", OUT)
