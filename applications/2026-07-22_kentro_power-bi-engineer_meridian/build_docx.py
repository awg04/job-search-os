#!/usr/bin/env python3
"""Build an ATS-friendly resume.docx for the Kentro Power BI Engineer (VA ESOM) role. MERIDIAN."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume.docx")
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(10)
for section in doc.sections:
    section.top_margin = section.bottom_margin = Pt(20)
    section.left_margin = section.right_margin = Pt(42)

def center(text, size=10.5, bold=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); p.paragraph_format.space_after = Pt(2); return p

def name_line(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18); p.paragraph_format.space_after = Pt(0)

def section_head(text):
    p = doc.add_paragraph(); r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x1F,0x38,0x64)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'1F3864')
    pbdr.append(bottom); pPr.append(pbdr)

def role(title, org_dates):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
    p.add_run(title).bold = True; p.runs[0].font.size = Pt(10)
    p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(1); r2 = p2.add_run(org_dates); r2.italic = True; r2.font.size = Pt(9.5)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(1)
    for i, seg in enumerate(text.split("**")):
        run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(9.5)

def para(text, size=9.5):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    for i, seg in enumerate(text.split("**")):
        run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(size)

name_line("Andrew Green")
center("Power BI Engineer — Power BI Desktop & Service · DAX · Power Query · SQL · Star Schema", 10.5, bold=True)
center("Fort Lauderdale, FL (Remote, ET)  |  561.339.1806  |  andrewgreen04@gmail.com  |  linkedin.com/in/agreen8", 9.5)

section_head("Summary")
para("Power BI Engineer with 10+ years in business intelligence on the Microsoft data stack — Power BI Desktop and Power BI Service, DAX, Power Query, SQL, and star-schema data modeling. Built three enterprise analytics platforms end to end across the full SDLC, translating complex datasets into executive-level insights in large, complex enterprise environments. Known for data accuracy (99%), performance optimization (25% faster processing), and stakeholder engagement that lifted adoption 35%.")

section_head("Technical Skills")
for line in [
 "**Power BI:** Power BI Desktop, Power BI Service, dataset/semantic modeling, star schema design, DAX (calculation groups), Power Query / M, paginated reports (SSRS / Report Builder)",
 "**SQL & integration:** SQL / T-SQL, relational databases (Azure SQL, SQL Server), ETL (SSIS, Azure Synapse, Data Factory), SharePoint & Excel integration, Power Automate, Power Apps; Microsoft Fabric in production (Lakehouse, Direct Lake)",
 "**Governance & SDLC:** dataset management, workspace organization, PBIP source control (Git), requirements gathering, data profiling & reconciliation, UAT",
 "**Certifications (in progress):** PL-300 Power BI Data Analyst · DP-700 Fabric Data Engineer · DP-203 Azure Data Engineer",
]:
    bullet(line)

section_head("Experience")
role("Business Intelligence Developer — Power BI Specialist (Contract/Consultant)", "Carnival Cruise Line — Remote  |  Jan 2025–Present")
for b in [
 "Design, develop, and maintain interactive **Power BI dashboards, reports, and datasets** (**Power BI Desktop and Service**) for fleet HVAC energy across 79 ships and 8 brands — KPI strips, performance trends, operational metrics, sensor-level drill-through diagnostics.",
 "Analyze complex data requirements and build **efficient, scalable data models** — reusable **DAX** measures, five calculation groups, **Power Query** cleansing and transformation over live Azure SQL — for engineering, operations, and executive audiences.",
 "Integrate Power BI with enterprise systems — **Azure SQL**, IoT telemetry historians, cloud BMS, on-prem **SQL Server** — via three ETL pipelines orchestrated in Azure Synapse with incremental loads and scheduled refresh; 99% data accuracy.",
 "Maintain **governance** best practices — **dataset management, workspace organization**, PBIP Git source control — and **document data sources, transformation logic, and reporting standards** for auditability and maintainability.",
 "Conduct **requirements gathering** and stakeholder working sessions; **troubleshoot data quality, performance, and modeling issues** (25% faster processing) and **train business users** — lifting dashboard adoption 35%.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer (Power BI)", "Basic Fun (consumer products / toys — CPG) — Boca Raton, FL  |  Mar 2023–Jan 2025")
for b in [
 "Designed the company's first enterprise **data warehouse** end to end — ingestion, ETL, **star schema design**, semantic layer, Power BI — plus **Microsoft Fabric** lakehouses publishing governed **golden datasets**.",
 "Built executive dashboards spanning **KPIs, operational metrics, and financial data** — sales, budget, forecast, inventory, variance-to-plan — driving a 20% Performance-to-Plan improvement; automated **scheduled report distribution** via **Power Automate**.",
 "Built ETL pipelines ingesting third-party point-of-sale data (web portal, FTP, email, flat file) into an Azure data lake; **incremental refresh** replaced full loads, cutting run times.",
 "Partnered with Finance, Sales, Demand Planning, and Operations on requirements and KPI logic — **translating business questions into analytical solutions** as Power BI reports and DAX measures.",
]:
    bullet(b)

role("Senior Business Intelligence Engineer", "Twin-Star International (consumer products, 3-company portfolio) — Delray Beach, FL  |  Oct 2017–Mar 2023")
for b in [
 "Architected the company's first enterprise **data warehouse**: dimensional models, **star schema**, **SSIS** ETL packages, and Azure Synapse for analytical processing.",
 "Automated P&L, Balance Sheet, and Cash Flow **financial reporting** across three portfolio companies — saving Finance 20–25 hours/month.",
 "Integrated Power BI with **SharePoint and Excel** — Power Apps write-back posting user commentary to a SharePoint tracker from inside Power BI; **SSRS / Report Builder** paginated reports for Excel/CSV exports.",
]:
    bullet(b)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
for i, seg in enumerate("**Earlier:** Business Analyst, AutoNation (2016-2017) - advanced DAX measures and SQL-based sales & pricing analytics  |  Financial Advisor, Merrill Lynch (2009-2016) - co-managed $250M in assets.  **Education:** B.S., Finance & Economics, Florida State University.".split("**")):
    run = p.add_run(seg); run.bold = (i % 2 == 1); run.font.size = Pt(9.5)

doc.save(OUT); print("SAVED:", OUT)


